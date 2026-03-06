"""
modules/document_processor.py
Document parsing module for PDF, DOCX, TXT, and Markdown files.
"""
import io
import re
from pathlib import Path
from typing import Optional


class DocumentProcessor:
    """
    Extracts plain text from uploaded documents.
    Supports: PDF, DOCX, TXT, MD
    """

    # Maximum text length to process
    MAX_TEXT_LENGTH = 100_000

    def extract_text(self, filename: str, content: bytes) -> str:
        """Dispatch extraction based on file extension."""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(content)
        elif ext == ".docx":
            return self._extract_docx(content)
        elif ext in (".txt", ".md", ".markdown", ".rst"):
            return self._extract_text(content)
        else:
            # Attempt plain text as fallback
            try:
                return content.decode("utf-8", errors="ignore")[:self.MAX_TEXT_LENGTH]
            except Exception:
                raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                pages.append(page.get_text("text"))
            doc.close()
            return self._clean_text("\n\n".join(pages))
        except ImportError:
            # Fallback to pdfplumber
            return self._extract_pdf_plumber(content)
        except Exception as e:
            print(f"PyMuPDF error: {e}, trying pdfplumber")
            return self._extract_pdf_plumber(content)

    def _extract_pdf_plumber(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber (fallback)."""
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            return self._clean_text("\n\n".join(pages))
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {e}")

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            return self._clean_text("\n".join(paragraphs))
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {e}")

    def _extract_text(self, content: bytes) -> str:
        """Decode plain text / markdown."""
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return self._clean_text(content.decode(encoding))
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Normalize whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        # Remove null bytes
        text = text.replace("\x00", "")
        return text.strip()[:self.MAX_TEXT_LENGTH]

    def chunk_text(self, text: str, chunk_size: int = 2000, overlap: int = 200) -> list:
        """
        Split long text into overlapping chunks for LLM processing.
        Returns list of text chunks.
        """
        if len(text) <= chunk_size:
            return [text]
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            if end < len(text):
                # Try to split on sentence boundary
                last_period = text.rfind(".", start, end)
                if last_period > start + chunk_size // 2:
                    end = last_period + 1
            chunks.append(text[start:end])
            start = end - overlap
        return chunks
